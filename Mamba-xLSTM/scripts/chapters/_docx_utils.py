"""Helper utilities for writing dissertation chapters as DOCX files.

Targets the ITB template (template-disertasi_Mei2019.docx) styling:
- Body text: Times New Roman 12pt, justify, 1.5 line spacing.
- Page: A4, margins top/bottom/right 3 cm, left 4 cm.
- Heading levels follow ITB nested numbering (e.g. III.1, III.1.1).
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_SIZE = Pt(12)
HEADING_BAB_SIZE = Pt(14)
SUBHEADING_SIZE = Pt(12)
LINE_SPACING = 1.5
SINGLE_SPACING = 1.0


def new_document() -> DocumentType:
    """Create a fresh ITB-style A4 document with correct margins."""

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(3.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


def _apply_run_font(run, *, bold: bool = False, italic: bool = False, size: Pt | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = size or BODY_SIZE
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)


def _apply_run_monospace(run, *, size: Pt | None = None) -> None:
    """Style a run as inline code (monospace font, slightly smaller).

    Used for ``\u0060...\u0060`` markdown segments (file paths, identifiers,
    YAML keys). The font name is forced via ``rFonts`` so Word does not fall
    back to the document default at render time.
    """

    run.font.name = MONO_FONT
    run.font.size = size or Pt(10.5)
    run.bold = False
    run.italic = False
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), MONO_FONT)


def add_chapter_title(doc: DocumentType, number_label: str, title: str) -> None:
    """Add a chapter heading like 'Bab III' followed by the chapter title centered."""

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(12)
    p1.paragraph_format.line_spacing = SINGLE_SPACING
    r1 = p1.add_run(number_label)
    _apply_run_font(r1, bold=True, size=HEADING_BAB_SIZE)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(18)
    p2.paragraph_format.line_spacing = SINGLE_SPACING
    r2 = p2.add_run(title.upper())
    _apply_run_font(r2, bold=True, size=HEADING_BAB_SIZE)


def add_heading(doc: DocumentType, number: str, title: str, level: int = 1) -> None:
    """Add a numbered subsection heading (e.g. 'III.1   Sintesis Model')."""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.keep_with_next = True
    text = f"{number}   {title}"
    r = p.add_run(text)
    if level == 1:
        _apply_run_font(r, bold=True, size=Pt(13))
    elif level == 2:
        _apply_run_font(r, bold=True, italic=False, size=Pt(12))
    else:
        _apply_run_font(r, bold=True, italic=True, size=Pt(12))


# --------------------------------------------------------------------------- #
# Inline LaTeX-style math rendering for prose paragraphs.
#
# Authors write small math snippets in prose using ``$...$`` (e.g. ``$x_t$``,
# ``$\sigma$``, ``$d_{\mathrm{model}} = 32$``). The helpers below convert those
# segments into properly formatted Word runs: Unicode for symbols, italic for
# math identifiers, sub/superscript for ``_`` and ``^`` groups, and roman text
# inside ``\mathrm{...}`` / ``\text{...}``.
# --------------------------------------------------------------------------- #

_LATEX_UNICODE = {
    # Lowercase Greek
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\varepsilon": "ε", r"\zeta": "ζ", r"\eta": "η",
    r"\theta": "θ", r"\vartheta": "ϑ", r"\iota": "ι", r"\kappa": "κ",
    r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν", r"\xi": "ξ", r"\pi": "π",
    r"\rho": "ρ", r"\sigma": "σ", r"\tau": "τ", r"\upsilon": "υ",
    r"\phi": "φ", r"\varphi": "φ", r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    # Uppercase Greek
    r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
    r"\Xi": "Ξ", r"\Pi": "Π", r"\Sigma": "Σ", r"\Upsilon": "Υ",
    r"\Phi": "Φ", r"\Psi": "Ψ", r"\Omega": "Ω",
    # Operators and binary ops
    r"\odot": "⊙", r"\otimes": "⊗", r"\oplus": "⊕", r"\ominus": "⊖",
    r"\cdot": "·", r"\times": "×", r"\div": "÷", r"\ast": "∗", r"\star": "⋆",
    r"\pm": "±", r"\mp": "∓", r"\circ": "∘",
    # Arrows
    r"\to": "→", r"\rightarrow": "→", r"\leftarrow": "←",
    r"\Rightarrow": "⇒", r"\Leftarrow": "⇐",
    r"\leftrightarrow": "↔", r"\Leftrightarrow": "⇔", r"\mapsto": "↦",
    r"\uparrow": "↑", r"\downarrow": "↓",
    # Relations
    r"\le": "≤", r"\leq": "≤", r"\ge": "≥", r"\geq": "≥",
    r"\ne": "≠", r"\neq": "≠", r"\approx": "≈", r"\sim": "∼",
    r"\equiv": "≡", r"\propto": "∝", r"\ll": "≪", r"\gg": "≫",
    # Set theory
    r"\in": "∈", r"\notin": "∉", r"\subset": "⊂", r"\subseteq": "⊆",
    r"\supset": "⊃", r"\supseteq": "⊇", r"\cup": "∪", r"\cap": "∩",
    r"\emptyset": "∅", r"\setminus": "∖",
    r"\mathbb{R}": "ℝ", r"\mathbb{N}": "ℕ", r"\mathbb{Z}": "ℤ",
    r"\mathbb{Q}": "ℚ", r"\mathbb{C}": "ℂ", r"\mathbb{E}": "𝔼",
    # Calculus / misc
    r"\infty": "∞", r"\partial": "∂", r"\nabla": "∇",
    r"\sum": "Σ", r"\prod": "∏", r"\int": "∫", r"\oint": "∮",
    r"\forall": "∀", r"\exists": "∃",
    r"\ldots": "…", r"\cdots": "⋯", r"\dots": "…", r"\vdots": "⋮",
    r"\top": "⊤", r"\bot": "⊥", r"\angle": "∠", r"\perp": "⊥",
    # Common accented letters used in this dissertation
    r"\hat{y}": "ŷ", r"\hat{x}": "x̂",
    r"\bar{y}": "ȳ", r"\bar{x}": "x̄", r"\bar{A}": "Ā", r"\bar{B}": "B̄",
    r"\tilde{x}": "x̃", r"\tilde{y}": "ỹ",
    # Spaces and modifiers
    r"\quad": "  ", r"\qquad": "    ",
    r"\,": "\u202f", r"\;": "\u2009", r"\:": "\u2009", r"\!": "",
    r"\&": "&", r"\%": "%", r"\#": "#",
    # Sizing modifiers (drop; the underlying delimiter is kept by the parser)
    r"\left": "", r"\right": "",
    r"\bigl": "", r"\bigr": "", r"\Bigl": "", r"\Bigr": "",
    r"\biggl": "", r"\biggr": "", r"\Biggl": "", r"\Biggr": "",
    r"\big": "", r"\Big": "", r"\bigg": "", r"\Bigg": "",
}

# Apply longest keys first so e.g. \mathbb{R} matches before \mathbb.
_LATEX_KEYS_SORTED = sorted(_LATEX_UNICODE.keys(), key=lambda k: -len(k))

# Generic accent commands rewritten with Unicode combining characters so the
# resulting glyph (e.g. ``\hat{c}`` -> ``ĉ``) can flow as part of a regular run.
_LATEX_REGEX_SUBS = [
    (re.compile(r"\\hat\{([^{}])\}"), lambda m: m.group(1) + "\u0302"),
    (re.compile(r"\\bar\{([^{}])\}"), lambda m: m.group(1) + "\u0304"),
    (re.compile(r"\\tilde\{([^{}])\}"), lambda m: m.group(1) + "\u0303"),
    (re.compile(r"\\dot\{([^{}])\}"), lambda m: m.group(1) + "\u0307"),
    (re.compile(r"\\vec\{([^{}])\}"), lambda m: m.group(1) + "\u20d7"),
    (re.compile(r"\\frac\{([^{}]+)\}\{([^{}]+)\}"), r"\1/\2"),
    (re.compile(r"\\sqrt\{([^{}]+)\}"), r"√(\1)"),
]

# Private-use placeholders for literal braces ``\{`` / ``\}``. They survive the
# parser's normal brace-stripping and are translated back to ``{`` / ``}`` when
# a run is emitted.
_LIT_BRACE_OPEN = "\ue000"
_LIT_BRACE_CLOSE = "\ue001"


def _apply_latex_substitutions(s: str) -> str:
    s = s.replace(r"\{", _LIT_BRACE_OPEN).replace(r"\}", _LIT_BRACE_CLOSE)
    for k in _LATEX_KEYS_SORTED:
        if len(k) > 1 and k.startswith("\\") and k[1].isalpha():
            # Word-boundary aware replacement: ``\Delta`` matches before
            # ``\,`` or ``_`` but not when followed by another letter.
            pattern = re.escape(k) + r"(?![A-Za-z])"
            s = re.sub(pattern, lambda _m, v=_LATEX_UNICODE[k]: v, s)
        else:
            s = s.replace(k, _LATEX_UNICODE[k])
    for rx, rep in _LATEX_REGEX_SUBS:
        s = rx.sub(rep, s)
    return s


def _find_matching_brace(text: str, start: int) -> int:
    if start >= len(text) or text[start] != "{":
        return start
    depth = 1
    i = start + 1
    while i < len(text):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return len(text)


def _read_script_argument(text: str, i: int) -> tuple[str, int]:
    """Read the argument that follows ``_`` or ``^``: ``{group}`` or one token."""

    if i >= len(text):
        return "", 0
    if text[i] == "{":
        end = _find_matching_brace(text, i)
        return text[i + 1 : end], end - i + 1
    if text[i] == "\\":
        m = re.match(r"\\[A-Za-z]+", text[i:])
        if m:
            return m.group(0), len(m.group(0))
        return text[i : i + 2], 2
    return text[i], 1


def _add_math_run(
    paragraph,
    text: str,
    *,
    italic: bool = False,
    bold: bool = False,
    subscript: bool = False,
    superscript: bool = False,
) -> None:
    if not text:
        return
    text = text.replace(_LIT_BRACE_OPEN, "{").replace(_LIT_BRACE_CLOSE, "}")
    if not text:
        return
    run = paragraph.add_run(text)
    _apply_run_font(run, bold=bold, italic=italic)
    if subscript:
        run.font.subscript = True
    if superscript:
        run.font.superscript = True


_FONT_COMMANDS = {
    "mathrm": (False, False),
    "text": (False, False),
    "operatorname": (False, False),
    "mathsf": (False, False),
    "mathtt": (False, False),
    "mathit": (True, False),
    "mathbf": (False, True),
}


def _render_inline_math(
    paragraph,
    math_text: str,
    *,
    italic_default: bool = True,
    base_subscript: bool = False,
    base_superscript: bool = False,
    inherit_bold: bool = False,
) -> None:
    text = _apply_latex_substitutions(math_text)

    i = 0
    n = len(text)
    buf: list[str] = []
    buf_italic: bool | None = None
    buf_bold = inherit_bold

    def flush() -> None:
        nonlocal buf, buf_italic
        if buf:
            _add_math_run(
                paragraph,
                "".join(buf),
                italic=bool(buf_italic),
                bold=buf_bold,
                subscript=base_subscript,
                superscript=base_superscript,
            )
            buf = []
            buf_italic = None

    while i < n:
        c = text[i]
        if c == "\\":
            m = re.match(r"\\([A-Za-z]+)\{", text[i:])
            if m and m.group(1) in _FONT_COMMANDS:
                cmd = m.group(1)
                start = i + len(m.group(0)) - 1
                end = _find_matching_brace(text, start)
                inner = text[start + 1 : end]
                flush()
                inner_italic, inner_bold = _FONT_COMMANDS[cmd]
                _render_inline_math(
                    paragraph,
                    inner,
                    italic_default=inner_italic,
                    base_subscript=base_subscript,
                    base_superscript=base_superscript,
                    inherit_bold=inherit_bold or inner_bold,
                )
                i = end + 1
                continue
            m = re.match(r"\\[A-Za-z]+", text[i:])
            if m:
                # Unknown command after substitutions: drop it silently rather
                # than emit a literal backslash.
                i += len(m.group(0))
                continue
            i += 1
            continue
        if c == "_":
            flush()
            inner, advance = _read_script_argument(text, i + 1)
            i += 1 + advance
            _render_inline_math(
                paragraph,
                inner,
                italic_default=italic_default,
                base_subscript=True,
                base_superscript=base_superscript,
                inherit_bold=inherit_bold,
            )
            continue
        if c == "^":
            flush()
            inner, advance = _read_script_argument(text, i + 1)
            i += 1 + advance
            _render_inline_math(
                paragraph,
                inner,
                italic_default=italic_default,
                base_subscript=base_subscript,
                base_superscript=True,
                inherit_bold=inherit_bold,
            )
            continue
        if c in "{}":
            i += 1
            continue
        if 0x0300 <= ord(c) <= 0x036F:
            # Combining diacritic: must inherit the previous char's italicness.
            if buf:
                buf.append(c)
            else:
                _add_math_run(
                    paragraph,
                    c,
                    italic=False,
                    bold=inherit_bold,
                    subscript=base_subscript,
                    superscript=base_superscript,
                )
            i += 1
            continue
        italic = italic_default and c.isalpha() and c.isascii()
        if buf and italic == buf_italic:
            buf.append(c)
        else:
            flush()
            buf = [c]
            buf_italic = italic
        i += 1

    flush()


def _split_inline_math(text: str) -> list[tuple[str, str]]:
    """Split ``text`` into a list of ``(kind, content)`` tokens.

    ``kind`` is ``"math"`` for a ``$...$`` segment (without the dollar signs)
    or ``"text"`` for the surrounding prose. Unmatched ``$`` is treated as
    plain prose.
    """

    out: list[tuple[str, str]] = []
    i = 0
    n = len(text)
    while i < n:
        if text[i] == "$":
            end = text.find("$", i + 1)
            if end == -1:
                out.append(("text", text[i:]))
                return out
            out.append(("math", text[i + 1 : end]))
            i = end + 1
            continue
        nxt = text.find("$", i)
        if nxt == -1:
            out.append(("text", text[i:]))
            return out
        out.append(("text", text[i:nxt]))
        i = nxt
    return out


def _render_markdown_segment(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Render a non-math segment with nested markdown emphasis.

    Supported syntax:

    * ``**bold**`` -> bold; recurses so inner ``*italic*`` / ``\u0060code\u0060`` still apply.
    * ``*italic*`` -> italic; recurses similarly.
    * ``\u0060code\u0060`` -> inline code (monospace, slightly smaller). Code spans are
      leaves: any surrounding bold/italic is intentionally not propagated so
      the snippet keeps its uniform code styling, matching common markdown
      renderers like GitHub.

    Unmatched opening tokens fall through to plain prose using the inherited
    bold/italic state.
    """

    i = 0
    n = len(text)
    while i < n:
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                _apply_run_font(paragraph.add_run(text[i:]), bold=bold, italic=italic)
                return
            _render_markdown_segment(
                paragraph, text[i + 2 : end], bold=True, italic=italic
            )
            i = end + 2
        elif text[i] == "*":
            end = text.find("*", i + 1)
            if end == -1:
                _apply_run_font(paragraph.add_run(text[i:]), bold=bold, italic=italic)
                return
            _render_markdown_segment(
                paragraph, text[i + 1 : end], bold=bold, italic=True
            )
            i = end + 1
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                _apply_run_font(paragraph.add_run(text[i:]), bold=bold, italic=italic)
                return
            r = paragraph.add_run(text[i + 1 : end])
            _apply_run_monospace(r)
            i = end + 1
        else:
            candidates = [
                idx
                for idx in (
                    text.find("**", i),
                    text.find("*", i),
                    text.find("`", i),
                )
                if idx != -1
            ]
            nxt = min(candidates) if candidates else -1
            if nxt == -1:
                r = paragraph.add_run(text[i:])
                _apply_run_font(r, bold=bold, italic=italic)
                return
            r = paragraph.add_run(text[i:nxt])
            _apply_run_font(r, bold=bold, italic=italic)
            i = nxt


def _add_runs_with_emphasis(paragraph, text: str) -> None:
    """Render text with markdown emphasis (``*italic*``, ``**bold**``) and
    inline LaTeX-style math (``$...$``).

    Math segments are converted to Unicode + Word run formatting so that
    ``$x_t$`` shows as italic *x* with a subscript italic *t*, ``$\\sigma$``
    shows as σ, ``$d_{\\mathrm{model}}$`` shows as italic *d* with a roman
    subscript ``model``, and so on. Outside math segments, the original
    markdown emphasis lexer applies.
    """

    for kind, segment in _split_inline_math(text):
        if not segment:
            continue
        if kind == "math":
            _render_inline_math(paragraph, segment)
        else:
            _render_markdown_segment(paragraph, segment)


def add_paragraph(
    doc: DocumentType,
    text: str,
    *,
    indent_first_line: bool = True,
    align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if indent_first_line:
        p.paragraph_format.first_line_indent = Cm(1.27)
    _add_runs_with_emphasis(p, text)


def add_bullets(doc: DocumentType, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        _add_runs_with_emphasis(p, item)


def add_numbered(doc: DocumentType, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        _add_runs_with_emphasis(p, item)


def add_equation(doc: DocumentType, text: str, label: str | None = None) -> None:
    """Render an equation as a centered text paragraph with optional label.

    Kept for backwards compatibility. For typeset math, prefer
    :func:`add_equation_image`, which renders LaTeX via matplotlib mathtext.
    """

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _apply_run_font(r)
    if label:
        r2 = p.add_run(f"        ({label})")
        _apply_run_font(r2)


_OMML_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ``mathml2omml==0.0.2`` has a known bug: for over-bar / under-bar / accent
# constructs (``\bar{X}``, ``\overline{X}``, ``\underline{X}``) it emits
# ``</m:groupChr>`` where it should emit ``</m:groupChrPr>``, producing
# malformed XML that lxml refuses to parse. The pattern that needs fixing is
# always ``<m:groupChrPr>`` containing only self-closing children followed by
# ``</m:groupChr>``. Rewrite that closing tag to the correct one.
_GROUP_CHR_PR_BUG = re.compile(
    r'(<m:groupChrPr>(?:<m:[A-Za-z]+\s+m:val="[^"]*"\s*/>)*)</m:groupChr>'
)


def _patch_mathml2omml_bugs(omml_str: str) -> str:
    return _GROUP_CHR_PR_BUG.sub(r"\1</m:groupChrPr>", omml_str)


def _latex_to_omml_element(latex: str):
    """Convert a LaTeX expression to a parsed ``<m:oMath>`` python-docx element.

    Returns ``None`` if the conversion pipeline fails (e.g. the LaTeX uses a
    construct that ``latex2mathml`` or ``mathml2omml`` does not support yet).
    """

    try:
        from latex2mathml.converter import convert as latex_to_mathml
        import mathml2omml
        from docx.oxml import parse_xml
    except ImportError:
        return None

    try:
        mathml = latex_to_mathml(latex)
        omml_str = mathml2omml.convert(mathml)
    except Exception:
        return None

    if not omml_str.startswith("<m:oMath"):
        return None

    omml_str = _patch_mathml2omml_bugs(omml_str)

    omml_str_with_ns = omml_str.replace(
        "<m:oMath>",
        f'<m:oMath xmlns:m="{_OMML_NAMESPACE}">',
        1,
    )
    try:
        return parse_xml(omml_str_with_ns)
    except Exception:
        return None


def add_equation_omml(
    doc: DocumentType,
    latex: str,
    *,
    label: str | None = None,
) -> bool:
    """Render a LaTeX expression as a native Word equation (OMML).

    The equation is editable in Word's Equation Editor and scales with the
    document font. Returns ``True`` on success and ``False`` if the LaTeX
    could not be converted (the caller can then fall back to PNG).
    """

    omath_element = _latex_to_omml_element(latex)
    if omath_element is None:
        return False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = SINGLE_SPACING
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p._p.append(omath_element)
    if label:
        run = p.add_run(f"\u2003\u2003\u2003\u2003({label})")
        _apply_run_font(run)
    return True


def add_equation_image(
    doc: DocumentType,
    latex: str,
    *,
    label: str | None = None,
    width_inches: float = 4.0,
    fontsize: int = 16,
    prefer_omml: bool = True,
) -> None:
    """Render a LaTeX equation centered.

    By default, the equation is emitted as a native Word OMML object so it is
    editable in Word's Equation Editor. If OMML conversion fails (unsupported
    LaTeX construct or missing dependency), falls back to a matplotlib-rendered
    PNG via :func:`render_equation`. Set ``prefer_omml=False`` to always use
    the PNG path.

    `latex` must be the math expression *without* enclosing ``$ ... $``.
    """

    if prefer_omml and add_equation_omml(doc, latex, label=label):
        return

    from chapters._equations import render_equation

    img_path = render_equation(latex, fontsize=fontsize)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = SINGLE_SPACING
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    if label:
        r2 = p.add_run(f"        ({label})")
        _apply_run_font(r2)


def add_diagram(
    doc: DocumentType,
    image_path: str | Path,
    caption: str,
    *,
    width_inches: float = 6.0,
) -> None:
    """Insert a diagram (architecture / flow / chart) followed by a caption.

    Convenience wrapper around :func:`add_figure` with a wider default size,
    suitable for the architectural diagrams generated by ``_diagrams.py``.
    """

    add_figure(doc, image_path, caption, width_inches=width_inches)


def _add_caption(
    doc: DocumentType,
    caption: str,
    *,
    label_word: str,
    space_before: int,
    space_after: int,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = SINGLE_SPACING
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    parts = caption.split(" ", 2)
    if len(parts) >= 2 and parts[0].lower() == label_word:
        r0 = p.add_run(f"{parts[0]} {parts[1]} ")
        _apply_run_font(r0, bold=True, size=Pt(11))
        rest = parts[2] if len(parts) >= 3 else ""
        if rest:
            _add_runs_with_emphasis(p, rest)
            for run in p.runs[1:]:
                run.font.size = Pt(10) if run.font.name == MONO_FONT else Pt(11)
    else:
        _add_runs_with_emphasis(p, caption)
        for run in p.runs:
            run.font.size = Pt(10) if run.font.name == MONO_FONT else Pt(11)


def add_figure_caption(doc: DocumentType, caption: str) -> None:
    _add_caption(doc, caption, label_word="gambar", space_before=2, space_after=12)


def add_table_caption(doc: DocumentType, caption: str) -> None:
    _add_caption(doc, caption, label_word="tabel", space_before=8, space_after=2)


def add_figure(doc: DocumentType, image_path: str | Path, caption: str, width_inches: float = 5.5) -> None:
    """Insert an image centered, then a numbered caption below."""

    image_path = Path(image_path)
    if not image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Gambar tidak ditemukan: {image_path}]")
        _apply_run_font(r, italic=True)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
    add_figure_caption(doc, caption)


def add_side_by_side_figure(
    doc: DocumentType,
    left_image: str | Path,
    right_image: str | Path,
    caption: str,
    *,
    image_width_inches: float = 2.7,
    left_label: str | None = None,
    right_label: str | None = None,
) -> None:
    """Place two images side-by-side using a 1x2 invisible table, then a single caption."""

    table = doc.add_table(rows=2 if (left_label or right_label) else 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    if left_label or right_label:
        hdr = table.rows[0].cells
        for cell, lbl in zip(hdr, (left_label or "", right_label or "")):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lbl)
            _apply_run_font(r, bold=True, size=Pt(10))
        img_row = table.rows[1].cells
    else:
        img_row = table.rows[0].cells

    for cell, path in zip(img_row, (left_image, right_image)):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = Path(path)
        if path.exists():
            p.add_run().add_picture(str(path), width=Inches(image_width_inches))
        else:
            r = p.add_run(f"[missing: {path.name}]")
            _apply_run_font(r, italic=True)

    add_figure_caption(doc, caption)


def add_table(
    doc: DocumentType,
    header: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    caption: str | None = None,
    col_widths: Sequence[float] | None = None,
    first_col_bold: bool = False,
) -> None:
    """Add a table with header and rows; optional caption is placed above per ITB convention."""

    if caption:
        add_table_caption(doc, caption)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = SINGLE_SPACING
        r = p.add_run(str(txt))
        _apply_run_font(r, bold=True, size=Pt(11))
        _shade_cell(cell, "D9E1F2")

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.line_spacing = SINGLE_SPACING
            _add_runs_with_emphasis(p, str(val))
            for run in p.runs:
                if run.font.name == MONO_FONT:
                    run.font.size = Pt(10)
                else:
                    run.font.size = Pt(11)
                if first_col_bold and c_idx == 0:
                    run.bold = True

    if col_widths is not None and len(col_widths) == len(header):
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


def add_page_break(doc: DocumentType) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_blockquote(doc: DocumentType, text: str) -> None:
    """Indented italicised note paragraph."""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _apply_run_font(r, italic=True)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def write_chapter(doc_factory, output_path: str | Path) -> dict:
    """Run a chapter-builder callable on a fresh document and save it.

    Returns a small dict with bookkeeping (path, paragraph count, table count).
    """

    doc = new_document()
    doc_factory(doc)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    n_tables = len(doc.tables)
    n_paragraphs = len(doc.paragraphs)
    n_images = sum(
        1
        for shape in doc.inline_shapes
    )
    return {
        "path": str(output_path),
        "paragraphs": n_paragraphs,
        "tables": n_tables,
        "inline_images": n_images,
    }
